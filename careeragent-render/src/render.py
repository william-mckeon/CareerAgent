#!/usr/bin/env python3
# ============================================================================
# careeragent-render - résumé markdown -> PDF/DOCX bytes (the substance)
# ============================================================================
#
# NO model, NO database, NO network. Pure, deterministic rendering. Given a
# résumé's markdown text and a target format, produce the document bytes.
#
# The one public entry point is a PURE FUNCTION (unit-testable without the API):
#
#   render(resume_text, fmt, title) -> bytes
#       Parse a focused résumé-markdown subset and lay it out as a clean,
#       professional PDF (reportlab) or DOCX (python-docx). Renders to an
#       in-memory BytesIO — never a temp file. Raises ValueError on bad input
#       (empty résumé / unsupported format) which backend.api maps to HTTP 400.
#
# Both renderers are pure-Python (reportlab + python-docx ship self-contained
# wheels — no system libraries, no apt layer). See specs/0001-render.md and
# docs/DATASHEET.md for the exact markdown subset supported and the honest
# limits: this is a focused résumé renderer, NOT a general markdown->PDF engine.
# ============================================================================

from __future__ import annotations

import io
import re
from typing import List, NamedTuple, Optional, Tuple

# reportlab — pure-Python PDF (platypus high-level layout).
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

# python-docx — pure-Python DOCX.
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SUPPORTED_FORMATS = ("pdf", "docx")

# ---------------------------------------------------------------------------
# Block model — the focused résumé-markdown subset we understand.
# ---------------------------------------------------------------------------
# A "block" is one laid-out unit. We deliberately support only what a résumé
# needs: headings (name / section / sub-section), bullet list items, paragraphs,
# and horizontal rules. Everything else degrades to a plain paragraph.
#
#   ("h1",     "<text>")   # `# ...`   — the name / top line (largest)
#   ("h2",     "<text>")   # `## ...`  — a section header (EXPERIENCE, SKILLS)
#   ("h3",     "<text>")   # `### ...` — a sub-header (a role / company line)
#   ("bullet", "<text>")   # `- ...` or `* ...`
#   ("para",   "<text>")   # a normal paragraph (blank-line separated)
#   ("hr",     "")         # `---` (3+ dashes on their own line)
#
# `<text>` still contains inline `**bold**` / `*italic*` markup; the renderers
# parse that per-block via _parse_inline.
Block = Tuple[str, str]

# A horizontal rule: 3+ dashes (only) on their own line. Checked BEFORE the
# bullet rule so `---` is a rule, not a `- ` bullet of "--".
_HR_RE = re.compile(r"^-{3,}$")

# Inline emphasis: `**bold**` (greedy-guarded, non-greedy inner) or `*italic*`.
# `**` is matched first (its alternative comes first) so bold wins over italic.
# Minimal by design — no nesting, no `__bold__`, no `_italic_`.
_INLINE_RE = re.compile(r"\*\*(?P<bold>.+?)\*\*|\*(?P<italic>.+?)\*")


# Per-laid-out-block character cap. reportlab's Paragraph line-breaking is ~O(n^2)
# in the block's length, so a single enormous paragraph (a résumé pasted as one
# long line) can take tens of seconds to MINUTES to lay out — the input-BYTE cap
# alone does NOT bound this (the cost is per-paragraph, not per-request). Splitting
# oversized blocks into bounded pieces keeps total layout ~linear (a 200 KB input
# becomes ~130 small blocks that render in well under a second). Normal résumé
# paragraphs/bullets are far under this, so real résumés are unchanged.
_MAX_BLOCK_CHARS = 1500
_MAX_HEADING_CHARS = 400

# Characters disallowed in XML 1.0 (C0 controls except \t, and lone surrogates).
# python-docx (lxml) RAISES on these, so an identical résumé could render as PDF
# but 500/400 as DOCX. Strip them once, up front, so both formats behave the same.
_BAD_XML_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\ud800-\udfff]")


def _sanitize_text(text: str) -> str:
    """Remove characters that XML/DOCX cannot represent (control chars, lone
    surrogates). Keeps \\t/\\n/\\r (newlines are normalized + split downstream)."""
    return _BAD_XML_CHARS.sub("", text or "")


def _chunk(text: str, limit: int) -> List[str]:
    """Split `text` into <=limit-char pieces at whitespace boundaries (hard-splits a
    single monster token). Always returns at least one piece."""
    text = text.strip()
    if len(text) <= limit:
        return [text]
    pieces: List[str] = []
    cur = ""
    for word in text.split(" "):
        while len(word) > limit:                 # a single word longer than the cap
            if cur:
                pieces.append(cur); cur = ""
            pieces.append(word[:limit]); word = word[limit:]
        if not cur:
            cur = word
        elif len(cur) + 1 + len(word) <= limit:
            cur += " " + word
        else:
            pieces.append(cur); cur = word
    if cur:
        pieces.append(cur)
    return pieces or [""]


def _bound_block_sizes(blocks: "List[Block]") -> "List[Block]":
    """Bound each laid-out block's length so no single reportlab Paragraph triggers
    the ~O(n^2) blow-up. Paragraphs/bullets are chunked; headings are truncated
    (a 1500-char heading is malformed anyway)."""
    out: List[Block] = []
    for kind, text in blocks:
        if kind in ("para", "bullet") and len(text) > _MAX_BLOCK_CHARS:
            out.extend((kind, piece) for piece in _chunk(text, _MAX_BLOCK_CHARS))
        elif kind in ("h1", "h2", "h3") and len(text) > _MAX_HEADING_CHARS:
            out.append((kind, text[:_MAX_HEADING_CHARS]))
        else:
            out.append((kind, text))
    return out


class InlineSpan(NamedTuple):
    text: str
    bold: bool
    italic: bool


def _parse_inline(text: str) -> List[InlineSpan]:
    """Split a line into (text, bold, italic) spans on `**bold**` / `*italic*`.

    A focused, non-nesting parser: each span carries at most one emphasis flag.
    Unmatched `*` / `**` stay literal. Always returns at least one span (which
    may carry empty text) so callers can iterate unconditionally.
    """
    spans: List[InlineSpan] = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            spans.append(InlineSpan(text[pos:m.start()], False, False))
        if m.group("bold") is not None:
            spans.append(InlineSpan(m.group("bold"), True, False))
        else:
            spans.append(InlineSpan(m.group("italic"), False, True))
        pos = m.end()
    if pos < len(text):
        spans.append(InlineSpan(text[pos:], False, False))
    if not spans:
        spans.append(InlineSpan(text, False, False))
    return spans


def _parse_blocks(resume_text: str) -> List[Block]:
    """Parse the résumé-markdown subset into an ordered list of blocks.

    Line-oriented and deterministic. Consecutive non-structural lines are joined
    into one paragraph; a blank line, heading, bullet, or rule flushes it.
    """
    blocks: List[Block] = []
    para_buf: List[str] = []

    def _flush_para() -> None:
        if para_buf:
            # Join wrapped lines of one paragraph with a single space.
            blocks.append(("para", " ".join(para_buf).strip()))
            para_buf.clear()

    # Normalize newlines so \r\n (Windows) and \r (old Mac) both split cleanly.
    for raw_line in resume_text.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            _flush_para()
            continue
        if _HR_RE.match(stripped):
            _flush_para()
            blocks.append(("hr", ""))
            continue
        if stripped.startswith("### "):
            _flush_para()
            blocks.append(("h3", stripped[4:].strip()))
            continue
        if stripped.startswith("## "):
            _flush_para()
            blocks.append(("h2", stripped[3:].strip()))
            continue
        if stripped.startswith("# "):
            _flush_para()
            blocks.append(("h1", stripped[2:].strip()))
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            _flush_para()
            blocks.append(("bullet", stripped[2:].strip()))
            continue
        # Anything else accumulates into the current paragraph.
        para_buf.append(stripped)

    _flush_para()
    # Bound per-block size so no single Paragraph triggers reportlab's O(n^2) cost.
    return _bound_block_sizes(blocks)


def _filename(fmt: str, title: Optional[str]) -> str:
    """Suggested download filename. Slugifies `title` if given, else "resume"."""
    base = "resume"
    if title and title.strip():
        slug = re.sub(r"[^a-z0-9]+", "-", title.strip().lower()).strip("-")
        if slug:
            base = slug[:60]  # keep filenames sane
    return f"{base}.{fmt}"


# ===========================================================================
# PDF (reportlab)
# ===========================================================================
def _escape_pdf(text: str) -> str:
    """Escape the characters reportlab's Paragraph treats as mini-HTML markup.

    reportlab Paragraph parses its text as a tiny XML/HTML dialect, so a raw
    `<`, `>`, or `&` in résumé text (e.g. "C++ & <legacy> systems") would break
    parsing or be silently dropped. Escape `&` FIRST (so we don't double-escape
    the `&` we introduce for `<`/`>`), then `<` and `>`.
    """
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_to_pdf_markup(text: str) -> str:
    """Turn one line into reportlab Paragraph markup: escape, then wrap emphasis.

    Escaping happens per-span on the RAW text, so the only `<`/`>` in the result
    are the `<b>`/`<i>` tags WE add — user `<...>` can never inject markup.
    """
    out: List[str] = []
    for span in _parse_inline(text):
        esc = _escape_pdf(span.text)
        if span.bold:
            esc = f"<b>{esc}</b>"
        if span.italic:
            esc = f"<i>{esc}</i>"
        out.append(esc)
    return "".join(out)


def _pdf_styles() -> dict:
    """Build the résumé paragraph styles (a clean, professional look).

    Name (H1) large; section headers (H2) bold + underscored by a rule; H3 a
    tight bold sub-header; body + bullets compact with tight leading.
    """
    ink = "#1a1a1a"
    accent = "#222222"
    return {
        "h1": ParagraphStyle(
            "resume_h1", fontName="Helvetica-Bold", fontSize=20, leading=23,
            textColor=ink, spaceBefore=0, spaceAfter=2, alignment=TA_LEFT,
        ),
        "h2": ParagraphStyle(
            "resume_h2", fontName="Helvetica-Bold", fontSize=12.5, leading=15,
            textColor=accent, spaceBefore=11, spaceAfter=3, alignment=TA_LEFT,
        ),
        "h3": ParagraphStyle(
            "resume_h3", fontName="Helvetica-Bold", fontSize=10.5, leading=13,
            textColor=ink, spaceBefore=6, spaceAfter=1, alignment=TA_LEFT,
        ),
        "para": ParagraphStyle(
            "resume_para", fontName="Helvetica", fontSize=10, leading=13,
            textColor=ink, spaceBefore=0, spaceAfter=4, alignment=TA_LEFT,
        ),
        "bullet": ParagraphStyle(
            "resume_bullet", fontName="Helvetica", fontSize=10, leading=13,
            textColor=ink, spaceBefore=0, spaceAfter=2, alignment=TA_LEFT,
            leftIndent=15, bulletIndent=3,
        ),
    }


def _render_pdf(blocks: List[Block], title: Optional[str]) -> bytes:
    """Lay the parsed blocks out as a PDF and return the bytes."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        title=(title or "Résumé"),
        author="",
        subject="",
        creator="careeragent-render",
    )
    styles = _pdf_styles()
    story: list = []

    for kind, text in blocks:
        if kind == "hr":
            story.append(Spacer(1, 2))
            story.append(HRFlowable(width="100%", thickness=0.6, color="#bbbbbb",
                                    spaceBefore=1, spaceAfter=5))
            continue
        markup = _inline_to_pdf_markup(text)
        if kind == "h1":
            story.append(Paragraph(markup, styles["h1"]))
            # A thin rule under the name gives a clean résumé header band.
            story.append(HRFlowable(width="100%", thickness=1.0, color="#444444",
                                    spaceBefore=3, spaceAfter=6))
        elif kind == "h2":
            story.append(Paragraph(markup, styles["h2"]))
            story.append(HRFlowable(width="100%", thickness=0.5, color="#cccccc",
                                    spaceBefore=0, spaceAfter=4))
        elif kind == "h3":
            story.append(Paragraph(markup, styles["h3"]))
        elif kind == "bullet":
            story.append(Paragraph(markup, styles["bullet"], bulletText="•"))
        else:  # para
            story.append(Paragraph(markup, styles["para"]))

    if not story:
        # A résumé that parsed to nothing still yields a valid (near-empty) PDF.
        story.append(Spacer(1, 1))

    doc.build(story)
    return buf.getvalue()


# ===========================================================================
# DOCX (python-docx)
# ===========================================================================
def _add_runs(paragraph, text: str) -> None:
    """Add bold/italic runs for one line to an existing python-docx paragraph."""
    for span in _parse_inline(text):
        run = paragraph.add_run(span.text)
        run.bold = span.bold
        run.italic = span.italic


def _add_hr_docx(document) -> None:
    """Append a horizontal rule as an empty paragraph with a bottom border.

    python-docx has no native rule, so we drop to the underlying OOXML and set
    a `w:pBdr/w:bottom` on an empty paragraph — the standard, portable recipe.
    """
    p = document.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _render_docx(blocks: List[Block], title: Optional[str]) -> bytes:
    """Lay the parsed blocks out as a DOCX and return the bytes."""
    document = Document()

    # Reasonable résumé margins.
    for section in document.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    if title and title.strip():
        document.core_properties.title = title.strip()

    for kind, text in blocks:
        if kind == "hr":
            _add_hr_docx(document)
        elif kind == "h1":
            _add_runs(document.add_heading("", level=0), text)      # Title style
        elif kind == "h2":
            _add_runs(document.add_heading("", level=1), text)      # Heading 1
        elif kind == "h3":
            _add_runs(document.add_heading("", level=2), text)      # Heading 2
        elif kind == "bullet":
            _add_runs(document.add_paragraph(style="List Bullet"), text)
        else:  # para
            _add_runs(document.add_paragraph(), text)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ===========================================================================
# Public entry point
# ===========================================================================
def render(resume_text: str, fmt: str, title: Optional[str] = None) -> bytes:
    """Render résumé markdown to PDF or DOCX bytes (pure, deterministic).

    Args:
      resume_text: the résumé as markdown text (see the supported subset above).
      fmt: "pdf" or "docx" (case-insensitive).
      title: optional document title / filename hint.

    Returns: the rendered document as bytes (in-memory; never a temp file).

    Raises:
      ValueError: empty/whitespace résumé, or an unsupported format. The API
      layer (backend.api) maps these to HTTP 400 with the contract messages.
    """
    if not resume_text or not resume_text.strip():
        raise ValueError("resume text is required to render.")

    normalized = (fmt or "").strip().lower()
    if normalized not in SUPPORTED_FORMATS:
        raise ValueError("format must be 'pdf' or 'docx'.")

    # Strip XML-illegal control chars up front so PDF and DOCX behave identically
    # (DOCX/lxml would otherwise raise on a stray control char that PDF tolerates).
    resume_text = _sanitize_text(resume_text)
    blocks = _parse_blocks(resume_text)
    if normalized == "pdf":
        return _render_pdf(blocks, title)
    return _render_docx(blocks, title)
