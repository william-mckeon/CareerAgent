#!/usr/bin/env python3
# ============================================================================
# careeragent-fetch - text extraction from an untrusted uploaded resume
# ============================================================================
#
# This module handles the FIRST untrusted files the system ingests. It never
# trusts the filename or the client-declared content-type — the document kind is
# decided by MAGIC BYTES, and every parser is fenced:
#   - PDF  : must start with %PDF. Parsed page-by-page with pdfplumber
#            (pdfminer.six, pure-Python), capped at MAX_PDF_PAGES. Parse errors →
#            400. Empty/whitespace text (scanned/image-only) → 422 (no OCR in P5).
#   - DOCX : must be a ZIP (PK\x03\x04) that CONTAINS word/document.xml. Guarded
#            against zip bombs (member count, uncompressed-size sum, expansion
#            ratio) BEFORE any member is decompressed, then parsed with
#            python-docx (its lxml parser sets resolve_entities=False, disabling
#            XXE external-entity expansion).
#   - .doc : legacy OLE (\xD0\xCF\x11\xE0) is explicitly rejected as unsupported.
#
# The actual bytes read are capped by the API layer (MAX_UPLOAD_BYTES) — we never
# rely on Content-Length alone.
# ============================================================================

from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass

import docx  # python-docx
import pdfplumber

# --- tunables ------------------------------------------------------------------
DEFAULT_MAX_PDF_PAGES = 30
DEFAULT_MAX_TEXT_CHARS = 100_000

# Zip-bomb guardrails for DOCX (a DOCX is a zip of XML; a hostile one can be tiny
# on disk yet expand to gigabytes). Checked from the central directory BEFORE any
# member is decompressed.
_MAX_ZIP_MEMBERS = 2_000
_MAX_DOCX_UNCOMPRESSED = 50_000_000    # 50 MB summed uncompressed size
_MAX_ZIP_RATIO = 200                   # uncompressed / compressed

# Magic bytes.
_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"
_OLE_MAGIC = b"\xd0\xcf\x11\xe0"       # legacy .doc / .xls (OLE2 compound file)

_SCANNED_PDF_MESSAGE = (
    "This PDF appears to be scanned/image-only; no text could be extracted "
    "(OCR is not supported)."
)


# --- exceptions: each carries the HTTP status the API layer should return ------
class ExtractProblem(Exception):
    """Base for everything /extract can go wrong with. status_code drives the HTTP."""
    status_code = 400

    def __init__(self, detail: str):
        self.detail = detail
        super().__init__(detail)


class UnsupportedFile(ExtractProblem):
    status_code = 415          # not a PDF/DOCX (or legacy .doc)


class FileTooLarge(ExtractProblem):
    status_code = 413          # read bytes exceeded MAX_UPLOAD_BYTES


class NoTextExtracted(ExtractProblem):
    status_code = 422          # scanned/image-only PDF with no text


class CorruptFile(ExtractProblem):
    status_code = 400          # malformed / unparseable document


@dataclass
class ExtractResult:
    text: str
    truncated: bool
    format: str
    chars: int


# ---------------------------------------------------------------------------
# magic-byte sniff
# ---------------------------------------------------------------------------
def _sniff(data: bytes) -> str:
    """Classify by leading bytes only. Returns 'pdf' | 'zip' | 'doc' | 'unknown'."""
    head = data[:4]
    if head == _PDF_MAGIC:
        return "pdf"
    if head == _OLE_MAGIC:
        return "doc"
    if head == _ZIP_MAGIC:
        return "zip"
    return "unknown"


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _extract_pdf(data: bytes, max_pages: int) -> str:
    parts: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    break
                parts.append(page.extract_text() or "")
    except Exception as err:  # pdfminer raises a zoo of exception types
        raise CorruptFile(f"could not parse PDF: {type(err).__name__}: {err}")
    return "\n".join(p for p in parts if p)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _extract_docx(data: bytes) -> str:
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as err:
        raise CorruptFile(f"not a valid DOCX (corrupt zip): {err}")

    # Guard against zip bombs from the central directory (no decompression yet),
    # and confirm this zip is actually a DOCX.
    with zf:
        infos = zf.infolist()
        if len(infos) > _MAX_ZIP_MEMBERS:
            raise CorruptFile("DOCX has too many entries (possible zip bomb)")
        if "word/document.xml" not in set(zf.namelist()):
            raise UnsupportedFile(
                "file is a zip but not a DOCX (missing word/document.xml)"
            )
        total_uncompressed = sum(i.file_size for i in infos)
        total_compressed = sum(i.compress_size for i in infos) or 1
        if total_uncompressed > _MAX_DOCX_UNCOMPRESSED:
            raise CorruptFile("DOCX uncompressed size too large (possible zip bomb)")
        if total_uncompressed / total_compressed > _MAX_ZIP_RATIO:
            raise CorruptFile("DOCX compression ratio too high (possible zip bomb)")

    # python-docx re-reads from a fresh stream; its lxml parser has
    # resolve_entities=False, so external-entity (XXE) expansion is disabled.
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as err:
        raise CorruptFile(f"could not parse DOCX: {type(err).__name__}: {err}")

    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(p for p in parts if p)


# ---------------------------------------------------------------------------
# dispatch
# ---------------------------------------------------------------------------
def extract_document(
    data: bytes,
    *,
    max_pdf_pages: int = DEFAULT_MAX_PDF_PAGES,
    max_text_chars: int = DEFAULT_MAX_TEXT_CHARS,
) -> ExtractResult:
    """Validate by magic bytes, extract text, and cap it. `data` is already
    byte-capped by the caller (the API layer enforces MAX_UPLOAD_BYTES)."""
    if not data:
        raise UnsupportedFile("empty upload")

    kind = _sniff(data)
    if kind == "pdf":
        fmt = "pdf"
        text = _extract_pdf(data, max_pdf_pages)
        if not text.strip():
            raise NoTextExtracted(_SCANNED_PDF_MESSAGE)
    elif kind == "zip":
        fmt = "docx"
        text = _extract_docx(data)
        if not text.strip():
            raise NoTextExtracted(
                "No text could be extracted from this DOCX (it may contain only images)."
            )
    elif kind == "doc":
        raise UnsupportedFile(
            "legacy .doc (OLE) files are not supported; please upload a PDF or DOCX"
        )
    else:
        raise UnsupportedFile(
            "unsupported file type; only PDF and DOCX resumes are supported"
        )

    truncated = False
    if len(text) > max_text_chars:
        text = text[:max_text_chars]
        truncated = True

    return ExtractResult(text=text, truncated=truncated, format=fmt, chars=len(text))
